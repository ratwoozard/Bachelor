# -*- coding: utf-8 -*-
"""Exam-ready Word export.

docs/SYNOPSIS_LUKA_FAERDIG.md + docs/figures/RESEARCH_STRUCTURE_SYNOPSIS_LUKA.png
    -> docs/SYNOPSIS_LUKA_FAERDIG.docx

Includes: cover page, automatic TOC, page numbers in footer, Times New Roman 12pt,
line spacing 1.5, 2.5 cm margins on all sides, figure embedded as PNG.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


ROOT = Path(__file__).resolve().parents[2]
MD_PATH = ROOT / "docs" / "SYNOPSIS_LUKA_FAERDIG.md"
DOCX_PATH = ROOT / "docs" / "SYNOPSIS_LUKA_FAERDIG.docx"
FIG_PNG = ROOT / "docs" / "figures" / "RESEARCH_STRUCTURE_SYNOPSIS_LUKA.png"

COVER = {
    "institution": "Københavns Erhvervsakademi (KEA)",
    "programme": "Professionsbachelor i IT og Økonomi",
    "module": "Theory of Science (videnskabsteori)",
    "type": "Synopsis — research design",
    "title": (
        "AI-baseret automatisering i bemandingsprocessen for indstilling af "
        "konsulenter hos Support Solutions ApS"
    ),
    "author": "Luka Christian Wigø",
    "student_id": "[indsæt studienummer]",
    "examiner": "Jens Rasmussen",
    "supervisor": "Saif",
    "deadline": "21. april 2026",
    "semester": "Forår 2026",
}


def _set_font(style, name: str = "Times New Roman", size: int | None = None) -> None:
    style.font.name = name
    if size is not None:
        style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_font(normal, "Times New Roman", 12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        style = doc.styles[name]
        _set_font(style, "Times New Roman", size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.3
        style.font.color.rgb = None  # inherit default

    title_style = doc.styles["Title"]
    _set_font(title_style, "Times New Roman", 22)


def _configure_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)


def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _configure_footer(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number_field(p)


def _add_toc_field(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Højreklik her og vælg Opdater felt — eller tryk F9 — for at generere "
        "indholdsfortegnelsen."
    )
    placeholder_run = OxmlElement("w:r")
    placeholder_run.append(placeholder)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder_run)
    run._r.append(end)


def _set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    tag = qn("w:updateFields")
    existing = settings.find(tag)
    if existing is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


def _add_cover_page(doc: Document) -> None:
    def _centered(text: str, size: int, *, bold: bool = False, italic: bool = False,
                  space_after: int = 6) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), "Times New Roman")

    for _ in range(4):
        doc.add_paragraph()

    _centered(COVER["institution"], 14, bold=True, space_after=2)
    _centered(COVER["programme"], 12)
    _centered(COVER["module"], 12, italic=True, space_after=24)

    _centered(COVER["type"], 13, italic=True, space_after=8)
    _centered(COVER["title"], 18, bold=True, space_after=48)

    _centered(COVER["author"], 13, space_after=2)
    _centered(f"Studienummer: {COVER['student_id']}", 11, italic=True, space_after=24)

    _centered(f"Eksaminator: {COVER['examiner']}", 12, space_after=2)
    _centered(f"Vejleder: {COVER['supervisor']}", 12, space_after=24)

    _centered(f"Afleveringsdato: {COVER['deadline']}", 12, space_after=2)
    _centered(COVER["semester"], 12, italic=True)

    # Hard page break to next section
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_toc_section(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Indholdsfortegnelse")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    _add_toc_field(doc)
    p2 = doc.add_paragraph()
    p2.add_run().add_break(WD_BREAK.PAGE)


def _add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _add_figure(doc: Document) -> None:
    if not FIG_PNG.is_file():
        p = doc.add_paragraph()
        p.add_run(f"[Figur mangler: {FIG_PNG.name}]").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIG_PNG), width=Cm(15))


def _add_body_from_markdown(doc: Document) -> None:
    raw = MD_PATH.read_text(encoding="utf-8")
    lines = raw.splitlines()

    # Skip everything until the first H2 ("## 1. Indledning") to avoid duplicating
    # the title block already shown on the cover page.
    start = next(
        (i for i, line in enumerate(lines) if line.startswith("## ")),
        0,
    )

    in_lit_list = False
    in_appendix = False

    for line in lines[start:]:
        s = line.rstrip()

        if s == "---":
            continue
        if not s:
            in_lit_list = False
            doc.add_paragraph()
            continue

        if s.startswith("## "):
            heading_text = s[3:]
            if heading_text.startswith("6. Litteraturliste"):
                in_lit_list = True
            if heading_text.startswith("7. Appendix"):
                in_appendix = True
                # Start appendix on new page
                doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
            doc.add_heading(heading_text, level=1)
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=2)
            continue

        if re.match(r"!\[([^\]]*)\]\(([^)]+)\)", s):
            _add_figure(doc)
            continue

        p = doc.add_paragraph()
        if in_lit_list:
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Cm(-1.5)
        _add_inline_runs(p, s)


def main() -> None:
    doc = Document()
    _configure_styles(doc)
    _configure_margins(doc)
    _configure_footer(doc)

    _add_cover_page(doc)
    _add_toc_section(doc)
    _add_body_from_markdown(doc)

    _set_update_fields_on_open(doc)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    print(f"Wrote {DOCX_PATH}")
    if not FIG_PNG.is_file():
        print(f"Warning: figure PNG missing at {FIG_PNG}")
    print("Manual step at first open in Word: right-click the TOC and choose")
    print("  'Update Field' (or press F9) to populate the indholdsfortegnelse.")
    print("  Also fill in student number on the cover page.")


if __name__ == "__main__":
    main()
