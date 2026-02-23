# gen_docx.py – Genererer Word-dokument fra introduktion med python-docx.
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from content_loader import load_introduction_md, get_plain_blocks, strip_md

def main():
    base = Path(__file__).resolve().parent
    out_dir = base.parent / "generated"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "01-introduction.docx"

    data = load_introduction_md()
    blocks = get_plain_blocks(data)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Titel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(data["title"])
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = None  # default
    doc.add_paragraph()

    doc.add_paragraph(strip_md(data["intro"]))
    doc.add_paragraph()

    for heading, para in blocks:
        if heading:
            p = doc.add_paragraph(heading)
            p.style = "Heading 2"
        if para:
            doc.add_paragraph(strip_md(para))

    doc.save(str(out_path))
    print(f"Skrevet: {out_path}")

if __name__ == "__main__":
    main()
